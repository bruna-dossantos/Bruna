#!/usr/bin/env python3
"""
render_extractions_doc.py  —  companion "Extraction Fields by Code" document

Mirrors the criteria doc, but for the extraction side: for each code → order type →
criterion, lists the extraction fields (directive, concept set, ICD candidates,
UMLS definition). Reads the extraction_fields.csv produced by
build_extraction_fields.py. Outputs Markdown and (optionally) .docx.

Usage:
  python3 render_extractions_doc.py extraction_fields.csv --out-md extractions_by_code.md
  .venv/bin/python render_extractions_doc.py extraction_fields.csv --out-docx extractions_by_code.docx
"""

import csv
import sys
import argparse
from pathlib import Path
from collections import OrderedDict


def load(csv_path):
    rows = list(csv.DictReader(open(csv_path)))
    # group: code -> order_type -> (criterion_n, criterion_title, type) -> [rows]
    tree = OrderedDict()
    for r in rows:
        code = f"{r['code']} — {r.get('modality','')}".strip(" —")
        ot = r.get("order_type") or "(single order type)"
        crit = (r.get("criterion_n", ""), r.get("criterion_title", ""), r.get("criterion_type", ""))
        tree.setdefault(code, OrderedDict()).setdefault(ot, OrderedDict()).setdefault(crit, []).append(r)
    return tree, len(rows)


def to_md(tree):
    out = ["# Extraction Fields by Code",
           "_Companion to the criteria doc. For each code → order type → criterion, "
           "the extraction fields Tennr needs: the go-find directive, the concept set "
           "(recall terms), ICD-10 candidates, and the UMLS definition._", ""]
    for code, ots in tree.items():
        out.append(f"## {code}")
        for ot, crits in ots.items():
            out.append(f"### Order Type: {ot}")
            for (n, title, ctype), fields in crits.items():
                out.append(f"**Criterion {n}. {title}**  ·  _{ctype}_")
                for f in fields:
                    out.append(f"- **{f['extraction_field']}** — {f['directive']}")
                    if f.get("concept_set"):
                        out.append(f"    - concept set: {f['concept_set']}")
                    if f.get("definition"):
                        out.append(f"    - definition ({f.get('definition_source','')}): {f['definition']}")
                    if f.get("icd10_candidates"):
                        out.append(f"    - ICD-10 candidates ({f.get('icd_source','')}): {f['icd10_candidates']}")
                out.append("")
    return "\n".join(out)


def to_docx(tree, out):
    from docx import Document
    from docx.shared import Pt, RGBColor
    d = Document()
    d.add_heading("Extraction Fields by Code", level=0)
    d.add_paragraph("Companion to the criteria doc. For each code → order type → "
                    "criterion, the extraction fields Tennr needs: the go-find directive, "
                    "the concept set (recall terms), ICD-10 candidates, and the UMLS "
                    "definition.").italic = True
    gray = RGBColor(0x66, 0x66, 0x66)
    for code, ots in tree.items():
        d.add_heading(code, level=1)
        for ot, crits in ots.items():
            d.add_heading(f"Order Type: {ot}", level=2)
            for (n, title, ctype), fields in crits.items():
                hp = d.add_paragraph()
                hp.add_run(f"Criterion {n}. {title}  ·  {ctype}").bold = True
                for f in fields:
                    fp = d.add_paragraph(style="List Bullet")
                    fp.add_run(f"{f['extraction_field']} — ").bold = True
                    fp.add_run(f["directive"])
                    for label, key, src in [("concept set", "concept_set", None),
                                            ("definition", "definition", "definition_source"),
                                            ("ICD-10", "icd10_candidates", "icd_source")]:
                        if f.get(key):
                            sp = d.add_paragraph(style="List Bullet 2")
                            tag = f" ({f.get(src,'')})" if src and f.get(src) else ""
                            r = sp.add_run(f"{label}{tag}: {f[key]}")
                            r.font.size = Pt(9)
                            r.font.color.rgb = gray
    d.save(out)


def main(argv=None):
    ap = argparse.ArgumentParser(description="Companion extraction-fields document")
    ap.add_argument("csv")
    ap.add_argument("--out-md")
    ap.add_argument("--out-docx")
    args = ap.parse_args(argv)
    if not args.out_md and not args.out_docx:
        ap.error("provide --out-md and/or --out-docx")
    tree, n = load(args.csv)
    if args.out_md:
        Path(args.out_md).write_text(to_md(tree))
        print(f"wrote {args.out_md}", file=sys.stderr)
    if args.out_docx:
        to_docx(tree, args.out_docx)
        print(f"wrote {args.out_docx}", file=sys.stderr)
    print(f"  {len(tree)} codes, {n} extraction-field rows", file=sys.stderr)


if __name__ == "__main__":
    main()
