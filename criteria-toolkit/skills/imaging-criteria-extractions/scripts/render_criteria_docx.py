#!/usr/bin/env python3
"""
render_criteria_docx.py  —  criteria .docx in the Tennr house template (SKILL Step 6)

Matches the format of the exemplar
`Criteria Updates/Criteria Docs/E0470 Nevada Medicaid — Checklist Format.docx`:
  - metadata table (Service Line, Plan Category, Policy Source, HCPCS Codes Covered)
  - Doc Criteria section
  - per code/order type: a colored SECTION BANNER, then each criterion as a badged
    block with a light header row, □ checklist sub-items, and an italic gray
    left-border policy-quote callout with its Source
  - "APPENDIX — Full Policy Language" on its own page

Requires python-docx (isolated venv if system Python is externally managed).

Usage:
  .venv/bin/python render_criteria_docx.py criteria.json --out criteria.docx
"""

import sys
import json
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK

sys.path.insert(0, str(Path(__file__).resolve().parent))
import criterion_view as V  # noqa: E402

BANNER = "2E5FA3"       # section banner fill
BANNER2 = "1F3864"      # exclusions / secondary segment
HEADER = "E8F0FB"       # criterion header row fill
AMBER = "FBF1DE"        # clinical-interpretation block fill
AMBER_INK = RGBColor(0xB7, 0x79, 0x1F)
POL_INK = RGBColor(0x2C, 0x5C, 0xC5)
GRAY = RGBColor(0x66, 0x66, 0x66)
BOX = "□"          # □

TYPE_LABEL = {"CLINICAL_INDICATION": "CLINICAL INDICATION", "PRIOR_WORKUP": "PRIOR WORKUP",
              "PRIOR_IMAGING": "PRIOR IMAGING", "METHODOLOGY": "METHODOLOGY", "CONTRAST": "CONTRAST",
              "THERAPY_LINKAGE": "THERAPY LINKAGE", "FREQUENCY": "FREQUENCY",
              "DOCUMENTATION": "DOCUMENTATION", "EXCLUSION": "EXCLUSION", "SPECIMEN": "SPECIMEN"}


def _shade(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "nil"); borders.append(e)
    tblPr.append(borders)


def banner(doc, text, fill=BANNER):
    t = doc.add_table(rows=1, cols=1); t.autofit = True; _no_borders(t)
    c = t.rows[0].cells[0]; _shade(c, fill)
    p = c.paragraphs[0]; r = p.add_run(text); r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(11)
    doc.add_paragraph()
    return t


def header_row(doc, badge, title, ctype):
    t = doc.add_table(rows=1, cols=1); t.autofit = True; _no_borders(t)
    c = t.rows[0].cells[0]; _shade(c, HEADER)
    p = c.paragraphs[0]
    b = p.add_run(f"{badge}   "); b.bold = True; b.font.size = Pt(11)
    r = p.add_run(title); r.bold = True; r.font.size = Pt(11)
    tag = p.add_run(f"    ·  {TYPE_LABEL.get(ctype, ctype)}")
    tag.font.size = Pt(8); tag.font.color.rgb = GRAY


def checklist(doc, definition):
    for line in [l for l in definition.split("\n") if l.strip()]:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(f"{BOX}  {line.strip()}")


def callout(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
    # left border
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    for k, v in (("w:val", "single"), ("w:sz", "18"), ("w:space", "8"), ("w:color", BANNER)):
        left.set(qn(k), v)
    pbdr.append(left); pPr.append(pbdr)
    r = p.add_run(text); r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(9)


def caption(doc, text, ink):
    """A small colored ALL-CAPS label to head a block."""
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(8); r.font.color.rgb = ink


def clinical_block(doc, od):
    """An amber-shaded block: our operational definition of a vague term.

    Visually distinct from the (blue) policy content so a reader can always tell
    'this is Tennr's interpretation, pending clinician sign-off' from 'this is the
    policy'.
    """
    t = doc.add_table(rows=1, cols=1); t.autofit = True; _no_borders(t)
    c = t.rows[0].cells[0]; _shade(c, AMBER)
    head = c.paragraphs[0]
    h = head.add_run(f'🩺 CLINICAL INTERPRETATION — how we read "{od.get("term","")}"')
    h.bold = True; h.font.size = Pt(8.5); h.font.color.rgb = AMBER_INK
    rev = V.reviewer_status(od)
    tag = head.add_run(f"   (reviewer {rev})"); tag.font.size = Pt(8); tag.italic = True; tag.font.color.rgb = GRAY
    for label, val in V.op_def_lines(od):
        p = c.add_paragraph()
        lr = p.add_run(f"{label}: "); lr.bold = True; lr.font.size = Pt(9); lr.font.color.rgb = AMBER_INK
        vr = p.add_run(val); vr.font.size = Pt(9)
    doc.add_paragraph()


def _codes(doc_json):
    groups = doc_json.get("groups") or [{"group_label": None, "codes": doc_json.get("codes", [])}]
    return [(g.get("group_label"), c) for g in groups for c in g["codes"]]


def build(doc_json, out):
    d = Document()
    p = doc_json.get("policy", {})
    title = d.add_heading(f"Qualification Criteria by Code — LCD {p.get('lcd','')}", level=0)
    d.add_paragraph(p.get("title", ""))

    # legend: what the two treatments mean (policy vs our interpretation)
    leg = d.add_paragraph()
    l1 = leg.add_run("How to read this: "); l1.bold = True; l1.font.size = Pt(9)
    l2 = leg.add_run("□ items and the blue Source line are the POLICY requirement. "); l2.font.size = Pt(9); l2.font.color.rgb = POL_INK
    l3 = leg.add_run("Amber 🩺 blocks are Tennr's CLINICAL INTERPRETATION of a vague term "
                     "(reviewer PENDING — a clinician signs off before go-live)."); l3.font.size = Pt(9); l3.font.color.rgb = AMBER_INK

    codes = [c["code"] for _, c in _codes(doc_json)]
    # metadata comes from the policy block so the doc is correct for ANY policy;
    # only the generic imaging fallbacks are hardcoded.
    plan_category = p.get("plan_category") or p.get("payer") or ("MEDICARE" if p.get("lcd") else "—")
    meta = [("Service Line", p.get("service_line") or "Imaging / Radiology"),
            ("Plan Category", plan_category),
            ("Policy Source", "; ".join(x for x in [f"LCD {p['lcd']}" if p.get("lcd") else "",
                                                      p.get("ncd_baseline", ""), p.get("article", "")] if x)),
            ("HCPCS Codes Covered", ", ".join(codes))]
    mt = d.add_table(rows=len(meta), cols=2); mt.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(meta):
        mt.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
        mt.rows[i].cells[1].text = v

    if p.get("doc_criteria"):
        d.add_heading("Doc Criteria", level=1)
        for dc in p["doc_criteria"]:
            d.add_heading(dc["document"], level=2)
            if dc.get("description"):
                d.add_paragraph(dc["description"]).runs[0].italic = True
            for f in dc.get("fields", []):
                d.add_paragraph(f, style="List Bullet")

    d.add_heading("Criteria by Code", level=1)
    last_group = None
    for group_label, c in _codes(doc_json):
        if group_label and group_label != last_group:
            d.add_heading(group_label, level=1); last_group = group_label
        d.add_heading(f"{c['code']} — {c.get('description','')}", level=2)
        m = " | ".join(x for x in [f"Modality: {c.get('modality','')}" if c.get("modality") else "",
                                   f"Contrast: {c.get('contrast','')}" if c.get("contrast") else ""] if x)
        if m:
            d.add_paragraph(m)
        ots = c.get("order_types") or [{"order_type": c.get("description", "Standard"),
                                        "criteria": c.get("criteria", []), "logic_expression": None}]
        for oi, ot in enumerate(ots):
            fill = BANNER2 if (ot.get("criteria") and all(cr.get("type") == "EXCLUSION"
                              for cr in ot["criteria"])) else BANNER
            label = ot.get("order_type", "Qualification")
            logic = ot.get("logic_expression")
            banner(d, f"{label}" + (f"  —  Qualifies when: {logic}" if logic else ""), fill)
            crits = ot.get("criteria", [])
            for ci, cr in enumerate(crits):
                header_row(d, str(cr.get("n", ci + 1)), cr.get("title", ""), cr.get("type", ""))
                policy_text, op_defs = V.split_criterion(cr)
                caption(d, "📄 POLICY REQUIREMENT", POL_INK)
                checklist(d, policy_text)
                if cr.get("source"):
                    callout(d, f"Source: {cr['source']}")
                for od in op_defs:
                    clinical_block(d, od)
                if ci < len(crits) - 1:
                    conn = d.add_paragraph(); rr = conn.add_run("AND"); rr.bold = True

    # Appendix — full policy language
    d.add_page_break()
    d.add_heading("APPENDIX — Full Policy Language", level=1)
    for group_label, c in _codes(doc_json):
        for ot in (c.get("order_types") or [{"order_type": "", "criteria": c.get("criteria", [])}]):
            for cr in ot.get("criteria", []):
                h = d.add_paragraph()
                h.add_run(f"{c['code']} · {ot.get('order_type','')} · "
                          f"{cr.get('n','')}. {cr.get('title','')}").bold = True
                policy_text, _ = V.split_criterion(cr)
                for line in policy_text.split("\n"):
                    if line.strip():
                        d.add_paragraph(line.strip())
                if cr.get("source"):
                    sp = d.add_paragraph(); sr = sp.add_run(f"Source: {cr['source']}")
                    sr.italic = True; sr.font.color.rgb = GRAY; sr.font.size = Pt(9)

    d.save(out)
    return len(codes)


def main(argv=None):
    ap = argparse.ArgumentParser(description="Render criteria JSON to Tennr-format .docx")
    ap.add_argument("criteria_json"); ap.add_argument("--out", required=True)
    args = ap.parse_args(argv)
    n = build(json.loads(Path(args.criteria_json).read_text()), args.out)
    print(f"wrote {args.out} ({n} codes, Tennr house format)", file=sys.stderr)


if __name__ == "__main__":
    main()
